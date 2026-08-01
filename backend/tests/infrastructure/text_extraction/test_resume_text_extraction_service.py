import asyncio
import io

import pymupdf
import pytest
from docx import Document

from application.exceptions import (
    DocumentTextExtractionError,
    NoExtractableTextError,
    UnsupportedDocumentTypeError,
)
from domain.entities.resume import ResumeMimeType
from infrastructure.text_extraction.resume_text_extraction_service import (
    ResumeTextExtractionService,
)


@pytest.fixture(autouse=True)
def run_thread_work_inline(monkeypatch: pytest.MonkeyPatch) -> None:
    async def run_inline(function, /, *args, **kwargs):
        return function(*args, **kwargs)

    monkeypatch.setattr(asyncio, "to_thread", run_inline)


def make_pdf(text: str | None = None) -> bytes:
    document = pymupdf.open()
    page = document.new_page()

    if text:
        page.insert_text((72, 72), text)

    content = document.tobytes()
    document.close()
    return content


def make_docx() -> bytes:
    document = Document()
    document.add_heading("Henok Asaye", level=1)
    document.add_paragraph("Backend Engineer")

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "Advanced"

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


@pytest.mark.asyncio
async def test_extracts_text_and_page_count_from_pdf() -> None:
    service = ResumeTextExtractionService()

    result = await service.extract(
        make_pdf("Henok Asaye\nBackend Engineer"),
        ResumeMimeType.PDF,
    )

    assert "Henok Asaye" in result.text
    assert "Backend Engineer" in result.text
    assert result.mime_type is ResumeMimeType.PDF
    assert result.page_count == 1
    assert result.character_count == len(result.text)


@pytest.mark.asyncio
async def test_extracts_paragraphs_and_tables_from_docx() -> None:
    service = ResumeTextExtractionService()

    result = await service.extract(
        make_docx(),
        ResumeMimeType.DOCX,
    )

    assert "Henok Asaye" in result.text
    assert "Backend Engineer" in result.text
    assert "Python | Advanced" in result.text
    assert result.mime_type is ResumeMimeType.DOCX
    assert result.page_count is None


@pytest.mark.asyncio
async def test_rejects_pdf_without_extractable_text() -> None:
    service = ResumeTextExtractionService()

    with pytest.raises(NoExtractableTextError):
        await service.extract(
            make_pdf(),
            ResumeMimeType.PDF,
        )


@pytest.mark.asyncio
async def test_rejects_corrupted_pdf() -> None:
    service = ResumeTextExtractionService()

    with pytest.raises(DocumentTextExtractionError):
        await service.extract(
            b"%PDF-invalid-content",
            ResumeMimeType.PDF,
        )


@pytest.mark.asyncio
async def test_rejects_unsupported_mime_type() -> None:
    service = ResumeTextExtractionService()

    with pytest.raises(UnsupportedDocumentTypeError):
        await service.extract(
            b"plain text",
            "text/plain",  # type: ignore[arg-type]
        )
