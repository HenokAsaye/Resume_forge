import asyncio
import io
from collections.abc import Iterable

import pymupdf
from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph

from application.exceptions import (
    DocumentTextExtractionError,
    NoExtractableTextError,
    UnsupportedDocumentTypeError,
)
from application.interfaces.services.document_text_extraction_service import (
    DocumentTextExtractionService,
    ExtractedDocument,
)
from domain.entities.resume import ResumeMimeType


class ResumeTextExtractionService(DocumentTextExtractionService):
    async def extract(
        self,
        content: bytes,
        mime_type: ResumeMimeType,
    ) -> ExtractedDocument:
        if not content:
            raise NoExtractableTextError("Document cannot be empty")

        try:
            normalized_mime_type = ResumeMimeType(mime_type)
        except (TypeError, ValueError) as exc:
            raise UnsupportedDocumentTypeError(
                f"Unsupported document MIME type: {mime_type}"
            ) from exc

        return await asyncio.to_thread(
            self._extract_sync, content, normalized_mime_type
        )

    def _extract_sync(
        self,
        content: bytes,
        mime_type: ResumeMimeType,
    ) -> ExtractedDocument:
        if mime_type is ResumeMimeType.PDF:
            return self._extract_pdf(content)
        if mime_type is ResumeMimeType.DOCX:
            return self._extract_docx(content)

        raise UnsupportedDocumentTypeError(
            f"No text extractor exists for MIME type: {mime_type}"
        )

    def _extract_pdf(self, content: bytes) -> ExtractedDocument:
        try:
            with pymupdf.open(stream=content, filetype="pdf") as document:
                if document.needs_pass:
                    raise DocumentTextExtractionError(
                        "Password-protected PDFs cannot be extracted"
                    )

                page_text = [page.get_text("text", sort=True) for page in document]
                page_count = document.page_count
        except DocumentTextExtractionError:
            raise
        except Exception as exc:
            raise DocumentTextExtractionError(
                "Unable to extract text from the PDF"
            ) from exc

        text = self._normalize_text("\n\n".join(page_text))
        self._ensure_text_exists(text)

        return ExtractedDocument(
            text=text,
            mime_type=ResumeMimeType.PDF,
            page_count=page_count,
        )

    def _extract_docx(self, content: bytes) -> ExtractedDocument:
        try:
            document = Document(io.BytesIO(content))
            text_parts = list(self._extract_docx_blocks(document.iter_inner_content()))
            text_parts.extend(self._extract_headers_and_footers(document))
        except Exception as exc:
            raise DocumentTextExtractionError(
                "Unable to extract text from the DOCX document"
            ) from exc

        text = self._normalize_text("\n\n".join(text_parts))
        self._ensure_text_exists(text)

        return ExtractedDocument(
            text=text,
            mime_type=ResumeMimeType.DOCX,
            page_count=None,
        )

    def _extract_docx_blocks(
        self,
        blocks: Iterable[Paragraph | Table],
    ) -> Iterable[str]:
        for block in blocks:
            if isinstance(block, Paragraph):
                if block.text.strip():
                    yield block.text
                continue

            if isinstance(block, Table):
                yield from self._extract_table(block)

    @staticmethod
    def _extract_table(table: Table) -> Iterable[str]:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]

            if cells:
                yield " | ".join(cells)

    def _extract_headers_and_footers(
        self,
        document: DocumentObject,
    ) -> list[str]:
        extracted_parts: list[str] = []
        seen_parts: set[str] = set()

        for section in document.sections:
            for container in (section.header, section.footer):
                part_name = str(container.part.partname)

                if part_name in seen_parts:
                    continue

                seen_parts.add(part_name)
                extracted_parts.extend(
                    self._extract_docx_blocks(container.iter_inner_content())
                )
        return extracted_parts

    @staticmethod
    def _normalize_text(text: str) -> str:
        normalized_lines: list[str] = []
        previous_line_was_empty = False

        for original_line in text.replace("\x00", "").splitlines():
            normalized_line = " ".join(original_line.split())

            if normalized_line:
                normalized_lines.append(normalized_line)
                previous_line_was_empty = False
                continue
            if normalized_lines and not previous_line_was_empty:
                normalized_lines.append("")
                previous_line_was_empty = True

        return "\n".join(normalized_lines).strip()

    @staticmethod
    def _ensure_text_exists(text: str) -> None:
        if not text:
            raise NoExtractableTextError("Document does not contain extractable text")
