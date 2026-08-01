import asyncio
import html
import io
import re

from application.dto.resume_schema import ResumeDocument
from application.interfaces.services.resume_export_service import (
    ExportedResume,
    ResumeExportFormat,
    ResumeExportService,
)
from docx import Document
from weasyprint import HTML


class ResumeDocumentExportService(ResumeExportService):
    async def export(
        self,
        resume: ResumeDocument,
        export_format: ResumeExportFormat,
        filename: str,
    ) -> ExportedResume:
        safe_filename = self._sanitize_filename(filename)
        return await asyncio.to_thread(
            self._export_sync,
            resume,
            export_format,
            safe_filename,
        )

    def _export_sync(
        self,
        resume: ResumeDocument,
        export_format: ResumeExportFormat,
        filename: str,
    ) -> ExportedResume:
        if export_format is ResumeExportFormat.DOCX:
            return ExportedResume(
                filename=f"{filename}.docx",
                media_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                content=self._build_docx(resume),
            )
        if export_format is ResumeExportFormat.PDF:
            return ExportedResume(
                filename=f"{filename}.pdf",
                media_type="application/pdf",
                content=self._build_pdf(resume),
            )
        raise ValueError(f"Unsupported export format: {export_format}")

    @staticmethod
    def _build_docx(resume: ResumeDocument) -> bytes:
        document = Document()
        document.add_heading(resume.contact.name, level=0)

        contact_parts = [
            resume.contact.email,
            resume.contact.phone,
            resume.contact.location,
            *resume.contact.links,
        ]
        document.add_paragraph(" | ".join(value for value in contact_parts if value))

        if resume.summary:
            document.add_heading("Professional Summary", level=1)
            document.add_paragraph(resume.summary)

        if resume.skills:
            document.add_heading("Skills", level=1)
            document.add_paragraph(", ".join(resume.skills))

        if resume.experience:
            document.add_heading("Experience", level=1)
            for experience in resume.experience:
                document.add_heading(
                    f"{experience.title} - {experience.company}",
                    level=2,
                )
                document.add_paragraph(f"{experience.start} - {experience.end}")
                for bullet in experience.bullets:
                    document.add_paragraph(bullet, style="List Bullet")

        if resume.education:
            document.add_heading("Education", level=1)
            for education in resume.education:
                document.add_heading(education.degree, level=2)
                document.add_paragraph(education.institution)
                document.add_paragraph(f"{education.start} - {education.end}")

        if resume.projects:
            document.add_heading("Projects", level=1)
            for project in resume.projects:
                document.add_heading(project.name, level=2)
                document.add_paragraph(project.description)
                if project.tech:
                    document.add_paragraph(f"Technologies: {', '.join(project.tech)}")

        if resume.certifications:
            document.add_heading("Certifications", level=1)
            for certification in resume.certifications:
                document.add_paragraph(
                    certification,
                    style="List Bullet",
                )

        stream = io.BytesIO()
        document.save(stream)
        return stream.getvalue()

    def _build_pdf(self, resume: ResumeDocument) -> bytes:
        return HTML(string=self._build_html(resume)).write_pdf()

    def _build_html(self, resume: ResumeDocument) -> str:
        def escaped(value: str) -> str:
            return html.escape(value)

        def list_items(values: list[str]) -> str:
            return "".join(f"<li>{escaped(value)}</li>" for value in values)

        experience_html = "".join(
            (
                f"<h3>{escaped(item.title)} - {escaped(item.company)}</h3>"
                f"<p class='date'>{escaped(item.start)} - "
                f"{escaped(item.end)}</p>"
                f"<ul>{list_items(item.bullets)}</ul>"
            )
            for item in resume.experience
        )
        education_html = "".join(
            (
                f"<h3>{escaped(item.degree)}</h3>"
                f"<p>{escaped(item.institution)}</p>"
                f"<p class='date'>{escaped(item.start)} - "
                f"{escaped(item.end)}</p>"
            )
            for item in resume.education
        )
        projects_html = "".join(
            (
                f"<h3>{escaped(item.name)}</h3>"
                f"<p>{escaped(item.description)}</p>"
                f"<p><strong>Technologies:</strong> "
                f"{escaped(', '.join(item.tech))}</p>"
            )
            for item in resume.projects
        )
        contact = " | ".join(
            escaped(value)
            for value in [
                resume.contact.email,
                resume.contact.phone,
                resume.contact.location,
                *resume.contact.links,
            ]
            if value
        )

        return f"""
        <!doctype html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                @page {{ size: A4; margin: 18mm; }}
                body {{
                    font-family: Arial, sans-serif;
                    color: #202020;
                    font-size: 10pt;
                    line-height: 1.45;
                }}
                h1 {{ margin-bottom: 4px; font-size: 24pt; }}
                h2 {{
                    margin-top: 18px;
                    margin-bottom: 6px;
                    border-bottom: 1px solid #555;
                    font-size: 13pt;
                }}
                h3 {{ margin-bottom: 2px; font-size: 11pt; }}
                p {{ margin: 3px 0; }}
                ul {{ margin-top: 4px; }}
                .contact {{ color: #444; }}
                .date {{ color: #555; font-style: italic; }}
            </style>
        </head>
        <body>
            <h1>{escaped(resume.contact.name)}</h1>
            <p class="contact">{contact}</p>
            <h2>Professional Summary</h2>
            <p>{escaped(resume.summary)}</p>
            <h2>Skills</h2>
            <p>{escaped(", ".join(resume.skills))}</p>
            <h2>Experience</h2>
            {experience_html}
            <h2>Education</h2>
            {education_html}
            <h2>Projects</h2>
            {projects_html}
            <h2>Certifications</h2>
            <ul>{list_items(resume.certifications)}</ul>
        </body>
        </html>
        """

    @staticmethod
    def _sanitize_filename(filename: str) -> str:
        normalized = re.sub(
            r"[^A-Za-z0-9._-]+",
            "_",
            filename.strip(),
        ).strip("._")
        return normalized or "optimized_resume"
